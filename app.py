
import gradio as gr
import pandas as pd
import difflib
from docx import Document
from docx.shared import RGBColor

def procesar_archivo(file):
    df = pd.read_excel(file)
    columna_1 = "Target"
    columna_2 = "Edited Target"

    if columna_1 not in df.columns or columna_2 not in df.columns:
        return None, "❌ Error: Columnas no encontradas en el archivo.", gr.update(visible=False)

    doc = Document()
    doc.add_heading('Diferencias en contexto', level=1)

    def añadir_diferencias_a_doc(doc, texto1, texto2):
        texto1 = str(texto1) if pd.notna(texto1) else ""
        texto2 = str(texto2) if pd.notna(texto2) else ""
        d = list(difflib.ndiff(texto1.split(), texto2.split()))
        parrafo = doc.add_paragraph()
        for palabra in d:
            if palabra.startswith("- "):
                run = parrafo.add_run(palabra[2:] + " ")
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.font.strike = True
            elif palabra.startswith("+ "):
                run = parrafo.add_run(palabra[2:] + " ")
                run.font.color.rgb = RGBColor(0, 0, 255)
            else:
                parrafo.add_run(palabra[2:] + " ")

    for index, row in df.iterrows():
        doc.add_paragraph(f"Fila {index + 2}:")
        añadir_diferencias_a_doc(doc, row[columna_1], row[columna_2])
        doc.add_paragraph()

    output_path = "/tmp/Correcciones_diferencias.docx"
    doc.save(output_path)

    return output_path, "✅ Archivo generado con éxito. Ya puedes descargarlo ⬇️", gr.update(visible=True)

theme = gr.themes.Default(primary_hue="blue")

with gr.Blocks(theme=theme) as iface:
    gr.Markdown("# 📝 Exportar correcciones de Excel a Word")
    gr.Markdown("Sube un archivo **Excel** con dos columnas: `Target` y `Edited Target`. El sistema generará un archivo **Word** resaltando las diferencias en rojo (eliminado) y azul (añadido).")

    with gr.Row():
        file_input = gr.File(label="📂 Subir archivo de Excel", type="filepath")

    message_output = gr.Textbox(label="Estado", interactive=False)
    file_output = gr.File(label="⬇️ Descargar archivo de Word", visible=False)

    submit_button = gr.Button("🔍 Procesar archivo")

    submit_button.click(procesar_archivo, inputs=file_input, outputs=[file_output, message_output, file_output])
        

iface.launch()

